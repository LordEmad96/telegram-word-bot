import telebot
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageEnhance, ImageFilter
import pytesseract
import io
import os
import re
import time
import hashlib
from datetime import datetime
import numpy as np
import cv2
from pdf2image import convert_from_bytes
import zipfile
import json
import threading
from collections import defaultdict
import logging
import gc

# ========== تنظیمات لاگینگ ==========
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# ========== تنظیمات ربات ==========
TOKEN = os.environ.get("TOKEN")
if not TOKEN:
    raise ValueError("توکن ربات در environment variables تنظیم نشده است!")

bot = telebot.TeleBot(TOKEN)

# تنظیمات پیشرفته با بهینه‌سازی برای Render
MAX_FILE_SIZE = 10 * 1024 * 1024  # کاهش به 10 مگابایت برای Render رایگان
SUPPORTED_IMAGE_FORMATS = ('.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.webp')
SUPPORTED_DOC_FORMATS = ('.pdf', '.zip')
OCR_LANGUAGES = 'fas+eng'
OCR_CONFIGS = {
    'fast': r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZآبپتثجچحخدذرزژسشصضطظعغفقکگلمنوهی',
    'accurate': r'--oem 3 --psm 6',
    'sparse': r'--oem 3 --psm 11',
    'single_line': r'--oem 3 --psm 7'
}

# کش برای پردازش‌های تکراری
processing_cache = {}
cache_lock = threading.Lock()

# آمار ربات
stats = defaultdict(int)
stats_lock = threading.Lock()

# ========== کلاس‌های پیشرفته ==========

class ImageProcessor:
    """پردازشگر پیشرفته تصاویر با بهینه‌سازی حافظه"""
    
    @staticmethod
    def remove_noise(image):
        """حذف نویز از تصویر"""
        try:
            if isinstance(image, Image.Image):
                image = np.array(image.convert('RGB'))
            
            # کاهش حجم برای پردازش سریع‌تر
            h, w = image.shape[:2]
            if h > 1000 or w > 1000:
                scale = 1000 / max(h, w)
                new_w = int(w * scale)
                new_h = int(h * scale)
                image = cv2.resize(image, (new_w, new_h))
            
            # اعمال فیلتر میانه
            denoised = cv2.medianBlur(image, 3)
            return denoised
        except Exception as e:
            logger.warning(f"خطا در حذف نویز: {e}")
            return image
    
    @staticmethod
    def correct_skew(image):
        """اصلاح کجی تصویر با محدودیت حافظه"""
        try:
            if isinstance(image, Image.Image):
                image = np.array(image.convert('L'))
            
            # کاهش حجم برای پردازش سریع‌تر
            h, w = image.shape[:2]
            if h > 800 or w > 800:
                scale = 800 / max(h, w)
                new_w = int(w * scale)
                new_h = int(h * scale)
                image = cv2.resize(image, (new_w, new_h))
            
            # تشخیص لبه‌ها
            edges = cv2.Canny(image, 50, 150, apertureSize=3)
            lines = cv2.HoughLines(edges, 1, np.pi/180, threshold=80)
            
            if lines is not None:
                angles = []
                for rho, theta in lines[:, 0]:
                    angle = np.degrees(theta) - 90
                    if abs(angle) < 45:
                        angles.append(angle)
                
                if angles and len(angles) > 0:
                    median_angle = np.median(angles)
                    if abs(median_angle) > 0.5:
                        (h, w) = image.shape[:2]
                        center = (w // 2, h // 2)
                        M = cv2.getRotationMatrix2D(center, median_angle, 1.0)
                        rotated = cv2.warpAffine(image, M, (w, h), 
                                                 flags=cv2.INTER_CUBIC,
                                                 borderMode=cv2.BORDER_REPLICATE)
                        return rotated
            return image
        except Exception as e:
            logger.warning(f"خطا در اصلاح کجی: {e}")
            return image
    
    @staticmethod
    def enhance_contrast(image):
        """بهبود کنتراست با محدودیت حافظه"""
        try:
            if isinstance(image, Image.Image):
                image = np.array(image.convert('L'))
            
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
            enhanced = clahe.apply(image)
            return enhanced
        except Exception as e:
            logger.warning(f"خطا در بهبود کنتراست: {e}")
            return image
    
    @staticmethod
    def sharpen_image(image):
        """تیز کردن تصویر"""
        try:
            if isinstance(image, Image.Image):
                image = np.array(image)
            
            kernel = np.array([[-1,-1,-1],
                               [-1, 9,-1],
                               [-1,-1,-1]])
            sharpened = cv2.filter2D(image, -1, kernel)
            return sharpened
        except Exception as e:
            logger.warning(f"خطا در تیز کردن: {e}")
            return image
    
    @staticmethod
    def auto_rotate(image):
        """تشخیص و چرخش خودکار تصویر بر اساس EXIF"""
        try:
            if hasattr(image, '_getexif'):
                exif = image._getexif()
                if exif:
                    orientation = exif.get(0x0112, 1)
                    rotate_map = {
                        3: Image.ROTATE_180,
                        6: Image.ROTATE_270,
                        8: Image.ROTATE_90
                    }
                    if orientation in rotate_map:
                        image = image.transpose(rotate_map[orientation])
        except Exception as e:
            logger.warning(f"خطا در چرخش خودکار: {e}")
        return image
    
    @staticmethod
    def preprocess_for_ocr(image):
        """پردازش کامل تصویر برای OCR با بهینه‌سازی حافظه"""
        try:
            # تبدیل به آرایه numpy
            if isinstance(image, Image.Image):
                img_array = np.array(image.convert('RGB'))
            else:
                img_array = image
            
            # محدودیت حجم
            h, w = img_array.shape[:2]
            if h > 1200 or w > 1200:
                scale = 1200 / max(h, w)
                new_w = int(w * scale)
                new_h = int(h * scale)
                img_array = cv2.resize(img_array, (new_w, new_h))
            
            # حذف نویز
            img_array = ImageProcessor.remove_noise(img_array)
            
            # تبدیل به خاکستری
            if len(img_array.shape) == 3:
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            else:
                gray = img_array
            
            # اصلاح کجی
            gray = ImageProcessor.correct_skew(gray)
            
            # بهبود کنتراست
            gray = ImageProcessor.enhance_contrast(gray)
            
            # تیز کردن
            gray = ImageProcessor.sharpen_image(gray)
            
            # آستانه‌گذاری تطبیقی
            binary = cv2.adaptiveThreshold(gray, 255, 
                                          cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                          cv2.THRESH_BINARY, 11, 2)
            
            return Image.fromarray(binary)
        except Exception as e:
            logger.error(f"خطا در پیش‌پردازش: {e}")
            return image

class TextFormatter:
    """فرمت‌دهنده پیشرفته متن خروجی"""
    
    @staticmethod
    def clean_text(text):
        """پاکسازی و نرمال‌سازی متن"""
        try:
            # حذف کاراکترهای اضافی
            text = re.sub(r'[^\w\s\u0600-\u06FF\uFB50-\uFDFF\uFE70-\uFEFF\.,!?;:()\[\]{}@#%&*+-=]', ' ', text)
            # حذف فاصله‌های اضافی
            text = re.sub(r'\s+', ' ', text)
            # حذف خطوط خالی
            text = '\n'.join(line.strip() for line in text.split('\n') if line.strip())
            return text.strip()
        except Exception as e:
            logger.warning(f"خطا در پاکسازی متن: {e}")
            return text
    
    @staticmethod
    def add_to_document(doc, text, title=None):
        """اضافه کردن متن به سند Word با فرمت‌دهی"""
        try:
            if title:
                heading = doc.add_heading(title, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = heading.runs[0]
                run.font.size = Pt(14)
                run.font.bold = True
            
            paragraphs = text.split('\n')
            for para_text in paragraphs[:50]:  # محدودیت پاراگراف برای جلوگیری از هنگ کردن
                if para_text.strip():
                    p = doc.add_paragraph()
                    run = p.add_run(para_text.strip()[:500])  # محدودیت طول هر پاراگراف
                    run.font.size = Pt(10)
                    
                    # تشخیص و تنظیم جهت متن
                    if re.search(r'[\u0600-\u06FF]', para_text):
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        run.font.rtl = True
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run.font.rtl = False
        except Exception as e:
            logger.error(f"خطا در افزودن به داکیومنت: {e}")

class DocumentProcessor:
    """پردازشگر اصلی اسناد با بهینه‌سازی برای Render"""
    
    @staticmethod
    def process_image(image_bytes, config_type='accurate'):
        """پردازش تصویر و استخراج متن"""
        try:
            img = Image.open(io.BytesIO(image_bytes))
            
            # محدودیت حجم تصویر
            if img.size[0] > 1500 or img.size[1] > 1500:
                img.thumbnail((1500, 1500), Image.Resampling.LANCZOS)
            
            # چرخش خودکار
            img = ImageProcessor.auto_rotate(img)
            
            # پیش‌پردازش
            processed_img = ImageProcessor.preprocess_for_ocr(img)
            
            # استخراج متن
            config = OCR_CONFIGS.get(config_type, OCR_CONFIGS['accurate'])
            text = pytesseract.image_to_string(processed_img, lang=OCR_LANGUAGES, config=config)
            
            # پاکسازی متن
            text = TextFormatter.clean_text(text)
            
            # پاکسازی حافظه
            del img
            del processed_img
            gc.collect()
            
            return text, len(text) > 50  # حداقل 50 کاراکتر معنی‌دار
        except Exception as e:
            logger.error(f"خطا در پردازش تصویر: {e}")
            return "", False
    
    @staticmethod
    def process_pdf(pdf_bytes, progress_callback=None):
        """پردازش PDF با مدیریت حافظه برای Render"""
        all_pages_text = []
        
        try:
            # کاهش کیفیت برای فایل‌های بزرگ (DPI 150 به جای 300)
            images = convert_from_bytes(pdf_bytes, dpi=150, fmt='jpeg')
            
            # محدود کردن تعداد صفحات
            max_pages = 15
            if len(images) > max_pages:
                images = images[:max_pages]
            
            for i, img in enumerate(images):
                if progress_callback:
                    progress_callback(i + 1, len(images))
                
                # کاهش حجم تصویر
                if img.size[0] > 1200 or img.size[1] > 1200:
                    img.thumbnail((1200, 1200), Image.Resampling.LANCZOS)
                
                # ذخیره موقت با کیفیت پایین‌تر
                img_byte_arr = io.BytesIO()
                img.save(img_byte_arr, format='JPEG', quality=60)
                
                # پردازش صفحه
                try:
                    text, success = DocumentProcessor.process_image(img_byte_arr.getvalue(), 'fast')  # استفاده از حالت fast
                    
                    if success and text.strip():
                        all_pages_text.append({
                            'page': i + 1,
                            'text': text[:3000],  # محدودیت متن هر صفحه
                            'has_text': True
                        })
                    else:
                        all_pages_text.append({
                            'page': i + 1,
                            'text': '[متنی در این صفحه تشخیص داده نشد]',
                            'has_text': False
                        })
                except Exception as page_error:
                    logger.error(f"خطا در صفحه {i+1}: {page_error}")
                    all_pages_text.append({
                        'page': i + 1,
                        'text': f'[خطا در پردازش صفحه {i+1}]',
                        'has_text': False
                    })
                
                # پاک کردن حافظه بعد از هر صفحه
                del img
                del img_byte_arr
                gc.collect()
            
            return all_pages_text
        except Exception as e:
            logger.error(f"خطا در پردازش PDF: {e}")
            return []
    
    @staticmethod
    def process_zip(zip_bytes, progress_callback=None):
        """پردازش فایل زیپ و استخراج متن از تمام فایل‌ها"""
        extracted_data = []
        
        try:
            with zipfile.ZipFile(io.BytesIO(zip_bytes)) as z:
                file_list = [f for f in z.namelist() 
                           if f.lower().endswith(SUPPORTED_IMAGE_FORMATS + SUPPORTED_DOC_FORMATS)]
                
                # محدود کردن تعداد فایل‌ها
                file_list = file_list[:10]
                
                for i, filename in enumerate(file_list):
                    if progress_callback:
                        progress_callback(i + 1, len(file_list))
                    
                    with z.open(filename) as f:
                        file_bytes = f.read()
                        
                        if filename.lower().endswith('.pdf') and len(file_bytes) < 5 * 1024 * 1024:  # حداکثر 5 مگابایت برای PDF داخل زیپ
                            pages = DocumentProcessor.process_pdf(file_bytes)
                            if pages:
                                extracted_data.append({
                                    'filename': filename,
                                    'type': 'pdf',
                                    'pages': pages[:5]  # حداکثر 5 صفحه
                                })
                        elif filename.lower().endswith(SUPPORTED_IMAGE_FORMATS):
                            text, success = DocumentProcessor.process_image(file_bytes, 'fast')
                            if success:
                                extracted_data.append({
                                    'filename': filename,
                                    'type': 'image',
                                    'text': text[:1000]  # محدودیت متن
                                })
                    
                    gc.collect()
            
            return extracted_data
        except Exception as e:
            logger.error(f"خطا در پردازش زیپ: {e}")
            return []

class TelegramBot:
    """کلاس اصلی ربات تلگرام"""
    
    def __init__(self):
        self.user_sessions = {}
        self.processing_tasks = {}
    
    def update_stats(self, file_type, success):
        """به‌روزرسانی آمار"""
        with stats_lock:
            stats['total_processed'] += 1
            stats[f'processed_{file_type}'] += 1
            if success:
                stats['successful'] += 1
            else:
                stats['failed'] += 1
    
    def send_progress(self, chat_id, current, total, message):
        """ارسال وضعیت پیشرفت"""
        try:
            percent = (current / total) * 100
            progress_bar = '█' * int(percent/10) + '░' * (10 - int(percent/10))
            text = f"{message}\n{progress_bar} {percent:.0f}% ({current}/{total})"
            bot.edit_message_text(text, chat_id, self.processing_tasks.get(chat_id))
        except:
            pass

bot_handler = TelegramBot()

# ========== هندلرهای ربات ==========

@bot.message_handler(commands=['start', 'help'])
def send_welcome(message):
    welcome_text = """
🤖 **ربات تبدیل فایل به Word**

📤 **فایل‌های پشتیبانی شده:**
• 📷 عکس‌ها (JPG, PNG, BMP)
• 📚 فایل‌های PDF (حداکثر 10 صفحه)
• 📦 فایل‌های ZIP (حداکثر 5 فایل)

⚠️ **توجه:**
• حداکثر حجم فایل: 10 مگابایت
• برای بهترین نتیجه، عکس با کیفیت بالا بفرستید

🚀 **فقط کافیست فایل خود را ارسال کنید!**
"""
    bot.reply_to(message, welcome_text, parse_mode='Markdown')

@bot.message_handler(commands=['stats'])
def show_stats(message):
    """نمایش آمار ربات"""
    stats_text = f"""
📊 **آمار ربات**

• کل فایل‌های پردازش شده: {stats.get('total_processed', 0)}
• پردازش‌های موفق: {stats.get('successful', 0)}
• پردازش‌های ناموفق: {stats.get('failed', 0)}

🏆 **نرخ موفقیت:** {((stats.get('successful', 0) / max(stats.get('total_processed', 1), 1)) * 100):.1f}%

🎯 **وضعیت:** آنلاین و فعال 🟢
"""
    bot.reply_to(message, stats_text, parse_mode='Markdown')

@bot.message_handler(content_types=['photo'])
def handle_photo(message):
    """پردازش عکس"""
    try:
        chat_id = message.chat.id
        msg = bot.reply_to(message, "📷 **در حال پردازش عکس...**", parse_mode='Markdown')
        bot_handler.processing_tasks[chat_id] = msg.message_id
        
        # دریافت عکس
        photo = message.photo[-1]
        file_info = bot.get_file(photo.file_id)
        downloaded = bot.download_file(file_info.file_path)
        
        # پردازش تصویر
        text, success = DocumentProcessor.process_image(downloaded, 'fast')
        
        if success and text:
            # ایجاد سند Word
            doc = Document()
            doc.add_heading('متن استخراج شده از تصویر', 0)
            doc.add_paragraph(f"تاریخ: {datetime.now().strftime('%Y/%m/%d - %H:%M:%S')}")
            doc.add_paragraph()
            
            # اضافه کردن متن
            TextFormatter.add_to_document(doc, text)
            
            # ذخیره و ارسال
            output_path = f"output_{chat_id}_{int(time.time())}.docx"
            doc.save(output_path)
            
            with open(output_path, 'rb') as f:
                bot.send_document(chat_id, f, caption="✅ **فایل Word ساخته شد!**")
            
            os.remove(output_path)
            bot.delete_message(chat_id, msg.message_id)
            
            bot_handler.update_stats('photo', True)
        else:
            bot.edit_message_text("❌ **متنی در این عکس تشخیص داده نشد**\nلطفاً عکس با کیفیت بهتری ارسال کنید.", 
                                chat_id, msg.message_id, parse_mode='Markdown')
            bot_handler.update_stats('photo', False)
    
    except Exception as e:
        logger.error(f"خطا در پردازش عکس: {e}")
        bot.reply_to(message, f"❌ **خطا:** {str(e)[:100]}")
        bot_handler.update_stats('photo', False)

@bot.message_handler(content_types=['document'])
def handle_document(message):
    """پردازش فایل‌های ضمیمه شده"""
    try:
        chat_id = message.chat.id
        file_name = message.document.file_name
        file_size = message.document.file_size
        
        # بررسی حجم فایل
        if file_size > MAX_FILE_SIZE:
            bot.reply_to(message, f"❌ **حجم فایل بیش از حد مجاز است!**\nحداکثر حجم: 10 مگابایت")
            return
        
        # بررسی فرمت فایل
        is_pdf = file_name.lower().endswith('.pdf')
        is_zip = file_name.lower().endswith('.zip')
        is_image = file_name.lower().endswith(SUPPORTED_IMAGE_FORMATS)
        
        if not (is_pdf or is_zip or is_image):
            bot.reply_to(message, f"❌ **فرمت فایل پشتیبانی نمی‌شود!**")
            return
        
        msg = bot.reply_to(message, f"📄 **در حال پردازش {file_name}...**", parse_mode='Markdown')
        bot_handler.processing_tasks[chat_id] = msg.message_id
        
        # دریافت فایل
        file_info = bot.get_file(message.document.file_id)
        downloaded = bot.download_file(file_info.file_path)
        
        doc = Document()
        doc.add_heading('متن استخراج شده از فایل', 0)
        doc.add_paragraph(f"نام فایل: {file_name}")
        doc.add_paragraph(f"تاریخ: {datetime.now().strftime('%Y/%m/%d - %H:%M:%S')}")
        doc.add_paragraph()
        
        if is_pdf:
            bot.edit_message_text(f"📄 **در حال پردازش PDF...**\nاین کار ممکن است چند لحظه طول بکشد.", 
                                chat_id, msg.message_id, parse_mode='Markdown')
            
            def progress_callback(current, total):
                if current % 2 == 0:  # به‌روزرسانی هر 2 صفحه
                    bot.edit_message_text(f"📄 **در حال پردازش PDF...**\nصفحه {current} از {total}", 
                                        chat_id, msg.message_id, parse_mode='Markdown')
            
            pages = DocumentProcessor.process_pdf(downloaded, progress_callback)
            
            if pages:
                for page in pages:
                    if page['has_text']:
                        doc.add_heading(f"صفحه {page['page']}", level=2)
                        TextFormatter.add_to_document(doc, page['text'])
                        doc.add_page_break()
                bot_handler.update_stats('pdf', True)
            else:
                doc.add_paragraph("❌ متنی در این PDF تشخیص داده نشد.")
                bot_handler.update_stats('pdf', False)
        
        elif is_zip:
            bot.edit_message_text(f"📦 **در حال پردازش فایل ZIP...**", 
                                chat_id, msg.message_id, parse_mode='Markdown')
            
            extracted_data = DocumentProcessor.process_zip(downloaded)
            
            if extracted_data:
                for data in extracted_data:
                    doc.add_heading(f"فایل: {data['filename']}", level=2)
                    
                    if data['type'] == 'pdf':
                        for page in data['pages']:
                            if page['has_text']:
                                doc.add_heading(f"صفحه {page['page']}", level=3)
                                TextFormatter.add_to_document(doc, page['text'])
                    else:
                        TextFormatter.add_to_document(doc, data['text'])
                    
                    doc.add_paragraph()
                
                bot_handler.update_stats('zip', True)
            else:
                doc.add_paragraph("❌ هیچ فایل پشتیبانی شده‌ای در ZIP یافت نشد.")
                bot_handler.update_stats('zip', False)
        
        elif is_image:
            text, success = DocumentProcessor.process_image(downloaded, 'fast')
            
            if success:
                TextFormatter.add_to_document(doc, text)
                bot_handler.update_stats('photo', True)
            else:
                doc.add_paragraph("❌ متنی در این عکس تشخیص داده نشد.")
                bot_handler.update_stats('photo', False)
        
        # ذخیره و ارسال فایل
        output_path = f"output_{chat_id}_{int(time.time())}.docx"
        doc.save(output_path)
        
        with open(output_path, 'rb') as f:
            bot.send_document(chat_id, f, caption="✅ **فایل Word با موفقیت ساخته شد!**")
        
        os.remove(output_path)
        bot.delete_message(chat_id, msg.message_id)
        bot.send_message(chat_id, "🎉 **پردازش انجام شد!**")
    
    except Exception as e:
        logger.error(f"خطا در پردازش فایل: {e}")
        bot.reply_to(message, f"❌ **خطا:** {str(e)[:150]}")
        bot_handler.update_stats('document', False)

@bot.message_handler(func=lambda message: True)
def handle_unknown(message):
    """پاسخ به پیام‌های ناشناخته"""
    bot.reply_to(message, 
                "🤔 **دستور ناشناخته!**\n\n"
                "لطفاً یک عکس یا فایل PDF ارسال کنید.\n"
                "برای راهنمایی بیشتر /start را بفرستید.",
                parse_mode='Markdown')

# ========== راه‌اندازی ربات ==========
if __name__ == '__main__':
    print("=" * 50)
    print("🤖 ربات تبدیل فایل به Word")
    print("📱 نسخه بهینه شده برای Render")
    print("=" * 50)
    print("✅ ربات با موفقیت راه‌اندازی شد!")
    print("=" * 50)
    
    try:
        bot.infinity_polling(timeout=60, long_polling_timeout=60)
    except KeyboardInterrupt:
        print("\n👋 ربات متوقف شد.")
    except Exception as e:
        print(f"❌ خطا در اجرای ربات: {e}")import telebot
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageEnhance, ImageFilter
import pytesseract
import io
import os
import re
import time
import hashlib
from datetime import datetime
import numpy as np
import cv2
from pdf2image import convert_from_bytes
import zipfile
import json
import threading
from collections import defaultdict
import logging
import gc

# ========== تنظیمات لاگینگ ==========
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# ========== تنظیمات ربات ==========
TOKEN = os.environ.get("TOKEN")
if not TOKEN:
    raise ValueError("توکن ربات در environment variables تنظیم نشده است!")

bot = telebot.TeleBot(TOKEN)

# تنظیمات پیشرفته با بهینه‌سازی برای Render
MAX_FILE_SIZE = 10 * 1024 * 1024  # کاهش به 10 مگابایت برای Render رایگان
SUPPORTED_IMAGE_FORMATS = ('.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.webp')
SUPPORTED_DOC_FORMATS = ('.pdf', '.zip')
OCR_LANGUAGES = 'fas+eng'
OCR_CONFIGS = {
    'fast': r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZآبپتثجچحخدذرزژسشصضطظعغفقکگلمنوهی',
    'accurate': r'--oem 3 --psm 6',
    'sparse': r'--oem 3 --psm 11',
    'single_line': r'--oem 3 --psm 7'
}

# کش برای پردازش‌های تکراری
processing_cache = {}
cache_lock = threading.Lock()

# آمار ربات
stats = defaultdict(int)
stats_lock = threading.Lock()

# ========== کلاس‌های پیشرفته ==========

class ImageProcessor:
    """پردازشگر پیشرفته تصاویر با بهینه‌سازی حافظه"""
    
    @staticmethod
    def remove_noise(image):
        """حذف نویز از تصویر"""
        try:
            if isinstance(image, Image.Image):
                image = np.array(image.convert('RGB'))
            
            # کاهش حجم برای پردازش سریع‌تر
            h, w = image.shape[:2]
            if h > 1000 or w > 1000:
                scale = 1000 / max(h, w)
                new_w = int(w * scale)
                new_h = int(h * scale)
                image = cv2.resize(image, (new_w, new_h))
            
            # اعمال فیلتر میانه
            denoised = cv2.medianBlur(image, 3)
            return denoised
        except Exception as e:
            logger.warning(f"خطا در حذف نویز: {e}")
            return image
    
    @staticmethod
    def correct_skew(image):
        """اصلاح کجی تصویر با محدودیت حافظه"""
        try:
            if isinstance(image, Image.Image):
                image = np.array(image.convert('L'))
            
            # کاهش حجم برای پردازش سریع‌تر
            h, w = image.shape[:2]
            if h > 800 or w > 800:
                scale = 800 / max(h, w)
                new_w = int(w * scale)
                new_h = int(h * scale)
                image = cv2.resize(image, (new_w, new_h))
            
            # تشخیص لبه‌ها
            edges = cv2.Canny(image, 50, 150, apertureSize=3)
            lines = cv2.HoughLines(edges, 1, np.pi/180, threshold=80)
            
            if lines is not None:
                angles = []
                for rho, theta in lines[:, 0]:
                    angle = np.degrees(theta) - 90
                    if abs(angle) < 45:
                        angles.append(angle)
                
                if angles and len(angles) > 0:
                    median_angle = np.median(angles)
                    if abs(median_angle) > 0.5:
                        (h, w) = image.shape[:2]
                        center = (w // 2, h // 2)
                        M = cv2.getRotationMatrix2D(center, median_angle, 1.0)
                        rotated = cv2.warpAffine(image, M, (w, h), 
                                                 flags=cv2.INTER_CUBIC,
                                                 borderMode=cv2.BORDER_REPLICATE)
                        return rotated
            return image
        except Exception as e:
            logger.warning(f"خطا در اصلاح کجی: {e}")
            return image
    
    @staticmethod
    def enhance_contrast(image):
        """بهبود کنتراست با محدودیت حافظه"""
        try:
            if isinstance(image, Image.Image):
                image = np.array(image.convert('L'))
            
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
            enhanced = clahe.apply(image)
            return enhanced
        except Exception as e:
            logger.warning(f"خطا در بهبود کنتراست: {e}")
            return image
    
    @staticmethod
    def sharpen_image(image):
        """تیز کردن تصویر"""
        try:
            if isinstance(image, Image.Image):
                image = np.array(image)
            
            kernel = np.array([[-1,-1,-1],
                               [-1, 9,-1],
                               [-1,-1,-1]])
            sharpened = cv2.filter2D(image, -1, kernel)
            return sharpened
        except Exception as e:
            logger.warning(f"خطا در تیز کردن: {e}")
            return image
    
    @staticmethod
    def auto_rotate(image):
        """تشخیص و چرخش خودکار تصویر بر اساس EXIF"""
        try:
            if hasattr(image, '_getexif'):
                exif = image._getexif()
                if exif:
                    orientation = exif.get(0x0112, 1)
                    rotate_map = {
                        3: Image.ROTATE_180,
                        6: Image.ROTATE_270,
                        8: Image.ROTATE_90
                    }
                    if orientation in rotate_map:
                        image = image.transpose(rotate_map[orientation])
        except Exception as e:
            logger.warning(f"خطا در چرخش خودکار: {e}")
        return image
    
    @staticmethod
    def preprocess_for_ocr(image):
        """پردازش کامل تصویر برای OCR با بهینه‌سازی حافظه"""
        try:
            # تبدیل به آرایه numpy
            if isinstance(image, Image.Image):
                img_array = np.array(image.convert('RGB'))
            else:
                img_array = image
            
            # محدودیت حجم
            h, w = img_array.shape[:2]
            if h > 1200 or w > 1200:
                scale = 1200 / max(h, w)
                new_w = int(w * scale)
                new_h = int(h * scale)
                img_array = cv2.resize(img_array, (new_w, new_h))
            
            # حذف نویز
            img_array = ImageProcessor.remove_noise(img_array)
            
            # تبدیل به خاکستری
            if len(img_array.shape) == 3:
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            else:
                gray = img_array
            
            # اصلاح کجی
            gray = ImageProcessor.correct_skew(gray)
            
            # بهبود کنتراست
            gray = ImageProcessor.enhance_contrast(gray)
            
            # تیز کردن
            gray = ImageProcessor.sharpen_image(gray)
            
            # آستانه‌گذاری تطبیقی
            binary = cv2.adaptiveThreshold(gray, 255, 
                                          cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                          cv2.THRESH_BINARY, 11, 2)
            
            return Image.fromarray(binary)
        except Exception as e:
            logger.error(f"خطا در پیش‌پردازش: {e}")
            return image

class TextFormatter:
    """فرمت‌دهنده پیشرفته متن خروجی"""
    
    @staticmethod
    def clean_text(text):
        """پاکسازی و نرمال‌سازی متن"""
        try:
            # حذف کاراکترهای اضافی
            text = re.sub(r'[^\w\s\u0600-\u06FF\uFB50-\uFDFF\uFE70-\uFEFF\.,!?;:()\[\]{}@#%&*+-=]', ' ', text)
            # حذف فاصله‌های اضافی
            text = re.sub(r'\s+', ' ', text)
            # حذف خطوط خالی
            text = '\n'.join(line.strip() for line in text.split('\n') if line.strip())
            return text.strip()
        except Exception as e:
            logger.warning(f"خطا در پاکسازی متن: {e}")
            return text
    
    @staticmethod
    def add_to_document(doc, text, title=None):
        """اضافه کردن متن به سند Word با فرمت‌دهی"""
        try:
            if title:
                heading = doc.add_heading(title, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = heading.runs[0]
                run.font.size = Pt(14)
                run.font.bold = True
            
            paragraphs = text.split('\n')
            for para_text in paragraphs[:50]:  # محدودیت پاراگراف برای جلوگیری از هنگ کردن
                if para_text.strip():
                    p = doc.add_paragraph()
                    run = p.add_run(para_text.strip()[:500])  # محدودیت طول هر پاراگراف
                    run.font.size = Pt(10)
                    
                    # تشخیص و تنظیم جهت متن
                    if re.search(r'[\u0600-\u06FF]', para_text):
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        run.font.rtl = True
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run.font.rtl = False
        except Exception as e:
            logger.error(f"خطا در افزودن به داکیومنت: {e}")

class DocumentProcessor:
    """پردازشگر اصلی اسناد با بهینه‌سازی برای Render"""
    
    @staticmethod
    def process_image(image_bytes, config_type='accurate'):
        """پردازش تصویر و استخراج متن"""
        try:
            img = Image.open(io.BytesIO(image_bytes))
            
            # محدودیت حجم تصویر
            if img.size[0] > 1500 or img.size[1] > 1500:
                img.thumbnail((1500, 1500), Image.Resampling.LANCZOS)
            
            # چرخش خودکار
            img = ImageProcessor.auto_rotate(img)
            
            # پیش‌پردازش
            processed_img = ImageProcessor.preprocess_for_ocr(img)
            
            # استخراج متن
            config = OCR_CONFIGS.get(config_type, OCR_CONFIGS['accurate'])
            text = pytesseract.image_to_string(processed_img, lang=OCR_LANGUAGES, config=config)
            
            # پاکسازی متن
            text = TextFormatter.clean_text(text)
            
            # پاکسازی حافظه
            del img
            del processed_img
            gc.collect()
            
            return text, len(text) > 50  # حداقل 50 کاراکتر معنی‌دار
        except Exception as e:
            logger.error(f"خطا در پردازش تصویر: {e}")
            return "", False
    
    @staticmethod
    def process_pdf(pdf_bytes, progress_callback=None):
        """پردازش PDF با مدیریت حافظه برای Render"""
        all_pages_text = []
        
        try:
            # کاهش کیفیت برای فایل‌های بزرگ (DPI 150 به جای 300)
            images = convert_from_bytes(pdf_bytes, dpi=150, fmt='jpeg')
            
            # محدود کردن تعداد صفحات
            max_pages = 15
            if len(images) > max_pages:
                images = images[:max_pages]
            
            for i, img in enumerate(images):
                if progress_callback:
                    progress_callback(i + 1, len(images))
                
                # کاهش حجم تصویر
                if img.size[0] > 1200 or img.size[1] > 1200:
                    img.thumbnail((1200, 1200), Image.Resampling.LANCZOS)
                
                # ذخیره موقت با کیفیت پایین‌تر
                img_byte_arr = io.BytesIO()
                img.save(img_byte_arr, format='JPEG', quality=60)
                
                # پردازش صفحه
                try:
                    text, success = DocumentProcessor.process_image(img_byte_arr.getvalue(), 'fast')  # استفاده از حالت fast
                    
                    if success and text.strip():
                        all_pages_text.append({
                            'page': i + 1,
                            'text': text[:3000],  # محدودیت متن هر صفحه
                            'has_text': True
                        })
                    else:
                        all_pages_text.append({
                            'page': i + 1,
                            'text': '[متنی در این صفحه تشخیص داده نشد]',
                            'has_text': False
                        })
                except Exception as page_error:
                    logger.error(f"خطا در صفحه {i+1}: {page_error}")
                    all_pages_text.append({
                        'page': i + 1,
                        'text': f'[خطا در پردازش صفحه {i+1}]',
                        'has_text': False
                    })
                
                # پاک کردن حافظه بعد از هر صفحه
                del img
                del img_byte_arr
                gc.collect()
            
            return all_pages_text
        except Exception as e:
            logger.error(f"خطا در پردازش PDF: {e}")
            return []
    
    @staticmethod
    def process_zip(zip_bytes, progress_callback=None):
        """پردازش فایل زیپ و استخراج متن از تمام فایل‌ها"""
        extracted_data = []
        
        try:
            with zipfile.ZipFile(io.BytesIO(zip_bytes)) as z:
                file_list = [f for f in z.namelist() 
                           if f.lower().endswith(SUPPORTED_IMAGE_FORMATS + SUPPORTED_DOC_FORMATS)]
                
                # محدود کردن تعداد فایل‌ها
                file_list = file_list[:10]
                
                for i, filename in enumerate(file_list):
                    if progress_callback:
                        progress_callback(i + 1, len(file_list))
                    
                    with z.open(filename) as f:
                        file_bytes = f.read()
                        
                        if filename.lower().endswith('.pdf') and len(file_bytes) < 5 * 1024 * 1024:  # حداکثر 5 مگابایت برای PDF داخل زیپ
                            pages = DocumentProcessor.process_pdf(file_bytes)
                            if pages:
                                extracted_data.append({
                                    'filename': filename,
                                    'type': 'pdf',
                                    'pages': pages[:5]  # حداکثر 5 صفحه
                                })
                        elif filename.lower().endswith(SUPPORTED_IMAGE_FORMATS):
                            text, success = DocumentProcessor.process_image(file_bytes, 'fast')
                            if success:
                                extracted_data.append({
                                    'filename': filename,
                                    'type': 'image',
                                    'text': text[:1000]  # محدودیت متن
                                })
                    
                    gc.collect()
            
            return extracted_data
        except Exception as e:
            logger.error(f"خطا در پردازش زیپ: {e}")
            return []

class TelegramBot:
    """کلاس اصلی ربات تلگرام"""
    
    def __init__(self):
        self.user_sessions = {}
        self.processing_tasks = {}
    
    def update_stats(self, file_type, success):
        """به‌روزرسانی آمار"""
        with stats_lock:
            stats['total_processed'] += 1
            stats[f'processed_{file_type}'] += 1
            if success:
                stats['successful'] += 1
            else:
                stats['failed'] += 1
    
    def send_progress(self, chat_id, current, total, message):
        """ارسال وضعیت پیشرفت"""
        try:
            percent = (current / total) * 100
            progress_bar = '█' * int(percent/10) + '░' * (10 - int(percent/10))
            text = f"{message}\n{progress_bar} {percent:.0f}% ({current}/{total})"
            bot.edit_message_text(text, chat_id, self.processing_tasks.get(chat_id))
        except:
            pass

bot_handler = TelegramBot()

# ========== هندلرهای ربات ==========

@bot.message_handler(commands=['start', 'help'])
def send_welcome(message):
    welcome_text = """
🤖 **ربات تبدیل فایل به Word**

📤 **فایل‌های پشتیبانی شده:**
• 📷 عکس‌ها (JPG, PNG, BMP)
• 📚 فایل‌های PDF (حداکثر 10 صفحه)
• 📦 فایل‌های ZIP (حداکثر 5 فایل)

⚠️ **توجه:**
• حداکثر حجم فایل: 10 مگابایت
• برای بهترین نتیجه، عکس با کیفیت بالا بفرستید

🚀 **فقط کافیست فایل خود را ارسال کنید!**
"""
    bot.reply_to(message, welcome_text, parse_mode='Markdown')

@bot.message_handler(commands=['stats'])
def show_stats(message):
    """نمایش آمار ربات"""
    stats_text = f"""
📊 **آمار ربات**

• کل فایل‌های پردازش شده: {stats.get('total_processed', 0)}
• پردازش‌های موفق: {stats.get('successful', 0)}
• پردازش‌های ناموفق: {stats.get('failed', 0)}

🏆 **نرخ موفقیت:** {((stats.get('successful', 0) / max(stats.get('total_processed', 1), 1)) * 100):.1f}%

🎯 **وضعیت:** آنلاین و فعال 🟢
"""
    bot.reply_to(message, stats_text, parse_mode='Markdown')

@bot.message_handler(content_types=['photo'])
def handle_photo(message):
    """پردازش عکس"""
    try:
        chat_id = message.chat.id
        msg = bot.reply_to(message, "📷 **در حال پردازش عکس...**", parse_mode='Markdown')
        bot_handler.processing_tasks[chat_id] = msg.message_id
        
        # دریافت عکس
        photo = message.photo[-1]
        file_info = bot.get_file(photo.file_id)
        downloaded = bot.download_file(file_info.file_path)
        
        # پردازش تصویر
        text, success = DocumentProcessor.process_image(downloaded, 'fast')
        
        if success and text:
            # ایجاد سند Word
            doc = Document()
            doc.add_heading('متن استخراج شده از تصویر', 0)
            doc.add_paragraph(f"تاریخ: {datetime.now().strftime('%Y/%m/%d - %H:%M:%S')}")
            doc.add_paragraph()
            
            # اضافه کردن متن
            TextFormatter.add_to_document(doc, text)
            
            # ذخیره و ارسال
            output_path = f"output_{chat_id}_{int(time.time())}.docx"
            doc.save(output_path)
            
            with open(output_path, 'rb') as f:
                bot.send_document(chat_id, f, caption="✅ **فایل Word ساخته شد!**")
            
            os.remove(output_path)
            bot.delete_message(chat_id, msg.message_id)
            
            bot_handler.update_stats('photo', True)
        else:
            bot.edit_message_text("❌ **متنی در این عکس تشخیص داده نشد**\nلطفاً عکس با کیفیت بهتری ارسال کنید.", 
                                chat_id, msg.message_id, parse_mode='Markdown')
            bot_handler.update_stats('photo', False)
    
    except Exception as e:
        logger.error(f"خطا در پردازش عکس: {e}")
        bot.reply_to(message, f"❌ **خطا:** {str(e)[:100]}")
        bot_handler.update_stats('photo', False)

@bot.message_handler(content_types=['document'])
def handle_document(message):
    """پردازش فایل‌های ضمیمه شده"""
    try:
        chat_id = message.chat.id
        file_name = message.document.file_name
        file_size = message.document.file_size
        
        # بررسی حجم فایل
        if file_size > MAX_FILE_SIZE:
            bot.reply_to(message, f"❌ **حجم فایل بیش از حد مجاز است!**\nحداکثر حجم: 10 مگابایت")
            return
        
        # بررسی فرمت فایل
        is_pdf = file_name.lower().endswith('.pdf')
        is_zip = file_name.lower().endswith('.zip')
        is_image = file_name.lower().endswith(SUPPORTED_IMAGE_FORMATS)
        
        if not (is_pdf or is_zip or is_image):
            bot.reply_to(message, f"❌ **فرمت فایل پشتیبانی نمی‌شود!**")
            return
        
        msg = bot.reply_to(message, f"📄 **در حال پردازش {file_name}...**", parse_mode='Markdown')
        bot_handler.processing_tasks[chat_id] = msg.message_id
        
        # دریافت فایل
        file_info = bot.get_file(message.document.file_id)
        downloaded = bot.download_file(file_info.file_path)
        
        doc = Document()
        doc.add_heading('متن استخراج شده از فایل', 0)
        doc.add_paragraph(f"نام فایل: {file_name}")
        doc.add_paragraph(f"تاریخ: {datetime.now().strftime('%Y/%m/%d - %H:%M:%S')}")
        doc.add_paragraph()
        
        if is_pdf:
            bot.edit_message_text(f"📄 **در حال پردازش PDF...**\nاین کار ممکن است چند لحظه طول بکشد.", 
                                chat_id, msg.message_id, parse_mode='Markdown')
            
            def progress_callback(current, total):
                if current % 2 == 0:  # به‌روزرسانی هر 2 صفحه
                    bot.edit_message_text(f"📄 **در حال پردازش PDF...**\nصفحه {current} از {total}", 
                                        chat_id, msg.message_id, parse_mode='Markdown')
            
            pages = DocumentProcessor.process_pdf(downloaded, progress_callback)
            
            if pages:
                for page in pages:
                    if page['has_text']:
                        doc.add_heading(f"صفحه {page['page']}", level=2)
                        TextFormatter.add_to_document(doc, page['text'])
                        doc.add_page_break()
                bot_handler.update_stats('pdf', True)
            else:
                doc.add_paragraph("❌ متنی در این PDF تشخیص داده نشد.")
                bot_handler.update_stats('pdf', False)
        
        elif is_zip:
            bot.edit_message_text(f"📦 **در حال پردازش فایل ZIP...**", 
                                chat_id, msg.message_id, parse_mode='Markdown')
            
            extracted_data = DocumentProcessor.process_zip(downloaded)
            
            if extracted_data:
                for data in extracted_data:
                    doc.add_heading(f"فایل: {data['filename']}", level=2)
                    
                    if data['type'] == 'pdf':
                        for page in data['pages']:
                            if page['has_text']:
                                doc.add_heading(f"صفحه {page['page']}", level=3)
                                TextFormatter.add_to_document(doc, page['text'])
                    else:
                        TextFormatter.add_to_document(doc, data['text'])
                    
                    doc.add_paragraph()
                
                bot_handler.update_stats('zip', True)
            else:
                doc.add_paragraph("❌ هیچ فایل پشتیبانی شده‌ای در ZIP یافت نشد.")
                bot_handler.update_stats('zip', False)
        
        elif is_image:
            text, success = DocumentProcessor.process_image(downloaded, 'fast')
            
            if success:
                TextFormatter.add_to_document(doc, text)
                bot_handler.update_stats('photo', True)
            else:
                doc.add_paragraph("❌ متنی در این عکس تشخیص داده نشد.")
                bot_handler.update_stats('photo', False)
        
        # ذخیره و ارسال فایل
        output_path = f"output_{chat_id}_{int(time.time())}.docx"
        doc.save(output_path)
        
        with open(output_path, 'rb') as f:
            bot.send_document(chat_id, f, caption="✅ **فایل Word با موفقیت ساخته شد!**")
        
        os.remove(output_path)
        bot.delete_message(chat_id, msg.message_id)
        bot.send_message(chat_id, "🎉 **پردازش انجام شد!**")
    
    except Exception as e:
        logger.error(f"خطا در پردازش فایل: {e}")
        bot.reply_to(message, f"❌ **خطا:** {str(e)[:150]}")
        bot_handler.update_stats('document', False)

@bot.message_handler(func=lambda message: True)
def handle_unknown(message):
    """پاسخ به پیام‌های ناشناخته"""
    bot.reply_to(message, 
                "🤔 **دستور ناشناخته!**\n\n"
                "لطفاً یک عکس یا فایل PDF ارسال کنید.\n"
                "برای راهنمایی بیشتر /start را بفرستید.",
                parse_mode='Markdown')

# ========== راه‌اندازی ربات ==========
if __name__ == '__main__':
    print("=" * 50)
    print("🤖 ربات تبدیل فایل به Word")
    print("📱 نسخه بهینه شده برای Render")
    print("=" * 50)
    print("✅ ربات با موفقیت راه‌اندازی شد!")
    print("=" * 50)
    
    try:
        bot.infinity_polling(timeout=60, long_polling_timeout=60)
    except KeyboardInterrupt:
        print("\n👋 ربات متوقف شد.")
    except Exception as e:
        print(f"❌ خطا در اجرای ربات: {e}")
